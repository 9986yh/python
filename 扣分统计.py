import openpyxl
from docx import Document

A3_53 = []
B3_53 = []
C3_53 = []
D3_53 = []
student_name = []
sheet_name = ['一连', '二连', '三连', '四连', '五连', '六连', '七连', '八连', '九连', '十连', '十一连']
for i in range(3, 60):
    B3_53.append('B' + str(i))
# print(B3_53)

for i in range(3, 60):
    D3_53.append('D' + str(i))
# print(D3_53)

for i in range(3, 60):
    A3_53.append('A' + str(i))
# print(A3_53)

for i in range(3, 60):
    C3_53.append('C' + str(i))
# print(C3_53)
work = '十二'
name = f'23大队第{work}周个人周统表.xlsx'
path = 'E:\\python学习\\excel\\大项目\\23大队第十二周个人扣分统计.docx'
doc = Document(path)

wb = openpyxl.load_workbook(name)
worksheet = wb.worksheets
# 获取所有sheet的文件名
for j in range(0, 11):
    sheet1 = worksheet[j]

    print(sheet_name[j])
    paragraph2 = doc.add_paragraph(f"{sheet_name[j]}")
    for i in range(len(C3_53)):
        # 扣分和通报次数

        if sheet1[C3_53[i]].value is not None and sheet1[D3_53[i]].value is not None:
            print(
                f"姓名：{sheet1[B3_53[i]].value}，{sheet1['C2'].value}:{sheet1[C3_53[i]].value}次，{sheet1['D2'].value}:{sheet1[D3_53[i]].value}次")
            paragraph1 = doc.add_paragraph(
                f"姓名：{sheet1[B3_53[i]].value}，{sheet1['C2'].value}:{sheet1[C3_53[i]].value}次，{sheet1['D2'].value}:{sheet1[D3_53[i]].value}次")
            doc.save(path)
            # 跳过
            continue
            # 结束
            # break
        # 扣分次数
        if sheet1[C3_53[i]].value is not None:
            print(f"姓名：{sheet1[B3_53[i]].value}，{sheet1['C2'].value}:{sheet1[C3_53[i]].value}次")
            paragraph1 = doc.add_paragraph(
                f"姓名：{sheet1[B3_53[i]].value}，{sheet1['C2'].value}:{sheet1[C3_53[i]].value}次")
            doc.save(path)
        # 通报次数
        if sheet1[D3_53[i]].value is not None:
            print(f"姓名：{sheet1[B3_53[i]].value}，{sheet1['D2'].value}:{sheet1[D3_53[i]].value}次")
            paragraph1 = doc.add_paragraph(
                f"姓名：{sheet1[B3_53[i]].value}，{sheet1['D2'].value}:{sheet1[D3_53[i]].value}次")
            doc.save(path)
print(f"{path}文件已生成")
# for i in range(len(C3_53)):
#     if sheet1[C3_53[i]].value is not None:
#         print(f"姓名：{sheet1[D3_53[i]].value}内务：{sheet1['D2'].value}{sheet1[D3_53[i]].value}次")

# 批量获取姓名
# for j in range(0, 11):
#     sheet1 = worksheet[j]
#     for i in range(len(B3_53)):
#         if sheet1[B3_53[i]].value is not None:
#             student_name.append(sheet1[B3_53[i]].value)
#     print(student_name)
#     student_name.clear()
